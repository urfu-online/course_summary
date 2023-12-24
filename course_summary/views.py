import logging
from datetime import timezone

import attr
from cms.djangoapps.contentstore.outlines import _make_section_data

from cms.djangoapps.contentstore.rest_api.v1.serializers import (
    CourseDetailsSerializer,
    CourseGradingModelSerializer,
    CourseGradingSerializer,
)
from cms.djangoapps.contentstore.utils import get_course_grading
from cms.djangoapps.coursegraph.tasks import coerce_types, serialize_item
from cms.djangoapps.models.settings.course_grading import CourseGradingModel
from common.djangoapps.util.views import ensure_valid_course_key
from django.conf import settings

from django.http import HttpResponse, HttpResponseForbidden
from django.shortcuts import render
from django.urls import reverse
from django.views.decorators.csrf import ensure_csrf_cookie

from docx import Document
from docx.shared import Pt
from lms.djangoapps.course_blocks.api import get_course_blocks
from lms.djangoapps.courseware.access import has_access
from lms.djangoapps.courseware.courses import (
    get_course,
    get_course_overview_with_access,
    get_course_with_access,
)
from opaque_keys import InvalidKeyError
from opaque_keys.edx.keys import CourseKey

from openedx.core.djangoapps.models.course_details import CourseDetails
from openedx.features.course_experience.utils import get_course_outline_block_tree
from xmodule.modulestore import ModuleStoreEnum
from xmodule.modulestore.django import modulestore

from .utils import get_course_outline_data, get_course_outline_data_dict


log = logging.getLogger(__name__)


@ensure_csrf_cookie
@ensure_valid_course_key
def summary(request, course_id):
    """
    Display the course's summary, or 404 if there is no such course.
    Assumes the course_id is in a valid format.
    """

    try:
        course_key = CourseKey.from_string(course_id)
    except InvalidKeyError:
        raise ValueError("Could not parse course_id {}".format(course_id))

    if not bool(has_access(request.user, "staff", get_course(course_key))):
        return HttpResponseForbidden("Forbidden")

    # course from modulestore
    course = get_course(course_key, depth=None)

    # CourseOverview; information that would be necessary to display a course
    course_overview = get_course_overview_with_access(request.user, "load", course_key)

    # CourseOutlineData
    course_outline_data = get_course_outline_data(course_key)

    # Extracted course information
    course_details = CourseDetails.fetch(course_key)

    # Model for CRUD operations pertaining to grading policy
    course_grading = CourseGradingModel.fetch(course_key)

    # root block of the course outline, with children as blocks including blocks with start dates in the future
    course_blocks = get_course_outline_block_tree(
        request, course_id, None, allow_start_dates_in_future=True
    )
    verticals = {}
    blocks = {}
    store = modulestore()

    for sequence in course_outline_data.sequences:
        stored_sequence = store.get_item(sequence, depth=None)
        verticals[sequence.block_id] = []
        for vertical in stored_sequence.get_children():
            verticals[sequence.block_id].append(vertical)
            # log.warning(dir(vertical))
            stored_vertical = store.get_item(vertical.location, depth=None)
            blocks[vertical] = []
            for block in stored_vertical.get_children():
                blocks[vertical].append(block)
                log.warning(
                    reverse(
                        "xblock_view", args=[course.id, block.location, "student_view"]
                    )
                )
                # log.warning(block.render('student_view'))
                # log.warning(block.student_view(_context=None))

    return render(
        request,
        "course_summary/base.html",
        {
            "course": course,
            "course_overview": course_overview,
            "course_outline_data": course_outline_data,
            "course_details": course_details,
            "course_grading": course_grading,
            # "course_blocks": course_blocks,
            "verticals": verticals,
            "blocks": blocks,
        },
    )


@ensure_csrf_cookie
@ensure_valid_course_key
def summary_docx(request, course_id):
    """
    Display the course's summary, or 404 if there is no such course.
    Assumes the course_id is in a valid format.
    """

    course_key = CourseKey.from_string(course_id)

    course = get_course_with_access(request.user, "load", course_key)
    if not bool(has_access(request.user, "staff", course)):
        return HttpResponseForbidden()

    store = modulestore()
    # course_outline_data = get_course_outline_data(course_id)

    with store.branch_setting(ModuleStoreEnum.Branch.published_only):
        course_blocks = get_course_outline_block_tree(request, course_id, None)

    document = Document()
    document.add_heading(course_blocks["display_name"], 0)

    section = document.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    run = header_p.add_run(
        f"{course_blocks['id']}\t\t{course_blocks['start'].strftime('%A, %d. %B %Y %I:%M%p')}"
    )
    font = run.font
    font.size = Pt(8)
    # header_p.style = document.styles["Signature"]

    # document.add_paragraph(str(attr.asdict(course_outline_data)))

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response[
        "Content-Disposition"
    ] = f"attachment; filename={course_blocks['display_name']}__summary.docx"
    document.save(response)

    return response
